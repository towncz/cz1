from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path.cwd()


def find_report():
    target = "报告1.docx".encode("unicode_escape").decode()
    for path in ROOT.glob("*.docx"):
        if path.name.encode("unicode_escape").decode() == target:
            return path
    raise FileNotFoundError("报告1.docx")


def set_font(run, size=11, bold=False, name="宋体"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    r = p.add_run(text)
    set_font(r, 13, True)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = Pt(22)
    r = p.add_run(text)
    set_font(r)


def add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"【图片占位：{text}】")
    set_font(r, 11, True)


def insert_after_paragraph(anchor, new_paragraphs):
    parent = anchor._p.getparent()
    index = parent.index(anchor._p)
    for offset, para in enumerate(new_paragraphs, start=1):
        parent.insert(index + offset, para._p)


path = find_report()
doc = Document(str(path))

full_text = "\n".join(p.text for p in doc.paragraphs)
if "Activity 模块用例图占位" not in full_text:
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "三、Activity 页面职责与跳转设计":
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("未找到插入位置")

    temp = Document()
    for style_name in ["Normal", "Heading 2"]:
        temp.styles[style_name].font.name = "宋体"
        temp.styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_h2(temp, "3.0 Activity 模块图示设计")
    add_para(temp, "为使页面关系和组件职责更加清晰，Activity 报告中预留以下图示位置。后续可根据实际代码绘制并插入对应图片，图中类名、组件名和流程名应与 MainActivity、DetailActivity、SearchActivity、AboutActivity、TicketQueryService、NetworkReceiver 等实际代码保持一致。")
    add_placeholder(temp, "Activity 模块用例图占位：用户浏览景点列表、查看详情、搜索地图、查看个人主页、返回列表")
    add_placeholder(temp, "Activity 模块组件图占位：MainActivity、DetailActivity、SearchActivity、AboutActivity 与 Service/Receiver 的协作关系")
    add_placeholder(temp, "Activity 模块类图占位：MainActivity、DetailActivity、SearchActivity、AboutActivity、Scenery、SceneryAdapter 的主要关系")
    add_placeholder(temp, "Activity 页面状态图占位：列表页、详情页、搜索页、个人主页之间的跳转和返回状态")

    insert_after_paragraph(anchor, list(temp.paragraphs))
    try:
        doc.save(str(path))
    except PermissionError:
        alt = ROOT / "报告1_activity_检查修正版.docx"
        doc.save(str(alt))
        print(f"SAVED_ALT={alt}")
        raise SystemExit(0)

print(f"UPDATED={path}")
